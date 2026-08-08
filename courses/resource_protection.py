"""Helpers for view-only Resources (PDF / Word / images) without download links."""

from __future__ import annotations

import mimetypes
import os
from html import escape
from pathlib import Path

from django.http import HttpResponse


IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}
PDF_EXTS = {".pdf"}
DOCX_EXTS = {".docx"}
DOC_EXTS = {".doc"}
TEXT_EXTS = {".txt", ".md", ".csv"}


def resource_file_ext(resource) -> str:
    name = ""
    if getattr(resource, "file", None) and resource.file:
        name = getattr(resource.file, "name", "") or ""
    return Path(name).suffix.lower()


def resource_preview_kind(resource) -> str:
    """Return: pdf | docx | doc | image | text | none | unknown"""
    if not getattr(resource, "file", None) or not resource.file:
        return "none"
    ext = resource_file_ext(resource)
    if ext in PDF_EXTS:
        return "pdf"
    if ext in DOCX_EXTS:
        return "docx"
    if ext in DOC_EXTS:
        return "doc"
    if ext in IMAGE_EXTS:
        return "image"
    if ext in TEXT_EXTS:
        return "text"
    return "unknown"


def guess_content_type(path_or_name: str) -> str:
    ctype, _ = mimetypes.guess_type(path_or_name)
    return ctype or "application/octet-stream"


def is_top_level_file_navigation(request) -> bool:
    """True when the browser is opening the file as a document (download/tab)."""
    fetch_dest = (request.headers.get("Sec-Fetch-Dest") or "").lower()
    fetch_mode = (request.headers.get("Sec-Fetch-Mode") or "").lower()

    # Embedded viewers / range streams
    if fetch_dest in (
        "iframe",
        "embed",
        "object",
        "image",
        "video",
        "audio",
        "empty",
    ):
        return False
    if fetch_mode in ("cors", "same-origin", "no-cors") and fetch_dest in ("", "empty"):
        return False

    if fetch_dest == "document" or fetch_mode == "navigate":
        return True

    if not fetch_dest and not fetch_mode and request.method == "GET":
        accept = (request.headers.get("Accept") or "").lower()
        if "text/html" in accept and "application/pdf" not in accept:
            return True
        # Explicit download intent
        if "attachment" in (request.headers.get("Content-Disposition") or "").lower():
            return True

    return False


def forbidden_download_response():
    return HttpResponse(
        "Downloading this resource is not allowed. "
        "Open it from the Resources page to view online only.",
        status=403,
        content_type="text/plain; charset=utf-8",
    )


def docx_to_protected_html(file_field) -> str:
    """Convert a .docx FileField to simple HTML paragraphs (escaped)."""
    try:
        from docx import Document
    except ImportError:
        return (
            "<p class='rd-preview-error'>Word preview is unavailable "
            "(python-docx not installed).</p>"
        )

    try:
        # Prefer filesystem path when available
        path = getattr(file_field, "path", None)
        if path and os.path.exists(path):
            document = Document(path)
        else:
            file_field.open("rb")
            try:
                document = Document(file_field)
            finally:
                file_field.close()
    except Exception as exc:
        return (
            f"<p class='rd-preview-error'>Could not open this Word document "
            f"for protected preview ({escape(str(exc)[:120])}).</p>"
        )

    parts: list[str] = []
    for para in document.paragraphs:
        text = (para.text or "").strip()
        if not text:
            parts.append("<div class='rd-para-spacer'></div>")
            continue
        style = (para.style.name if para.style else "") or ""
        tag = "p"
        cls = "rd-para"
        if style.startswith("Heading"):
            level = "".join(ch for ch in style if ch.isdigit()) or "2"
            tag = f"h{min(int(level), 4)}"
            cls = "rd-heading"
        parts.append(f"<{tag} class='{cls}'>{escape(text)}</{tag}>")

    for table in document.tables:
        parts.append("<table class='rd-table'>")
        for row in table.rows:
            parts.append("<tr>")
            for cell in row.cells:
                parts.append(f"<td>{escape((cell.text or '').strip())}</td>")
            parts.append("</tr>")
        parts.append("</table>")

    if not parts:
        return "<p class='rd-preview-error'>This document has no readable text.</p>"
    return "\n".join(parts)


def text_file_to_protected_html(file_field) -> str:
    try:
        file_field.open("rb")
        try:
            raw = file_field.read()
        finally:
            file_field.close()
        text = raw.decode("utf-8", errors="replace")
    except Exception as exc:
        return f"<p class='rd-preview-error'>Could not read file ({escape(str(exc)[:100])}).</p>"
    # Cap huge files
    if len(text) > 200_000:
        text = text[:200_000] + "\n\n… (preview truncated)"
    return f"<pre class='rd-pre'>{escape(text)}</pre>"
