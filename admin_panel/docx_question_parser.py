"""
Word (.doc / .docx) question-bank importer.

Pipeline (used by Admin → Upload Questions):

    DOC / DOCX Upload
            ↓
    Prepare / Normalize document   (OOXML namespace fix)
            ↓
    DOC → DOCX                     (LibreOffice when needed)
            ↓
    Extract equations              (OMML + WMF/EMF/PNG)
            ↓
    Equation → LaTeX               (pandoc / pure-Python fallback)
            ↓
    Parse questions + A/B/C/D
            ↓
    Parse answer sheet             (numbered answer tables)
            ↓
    Preview questions              (admin UI — not this module)
            ↓
    Edit if required               (admin UI)
            ↓
    Import Questions → database

Parses past-paper style Word documents shaped like:

    1. Question text ...
    (a) option text   (b) option text
    (c) option text   (d) option text
    ...
    Answer Sheet
    1  2  3  ...
    b  c  c  ...

Equations in these banks are almost always embedded as:
  - Native Word OMML math (m:oMath)  ← preferred
  - WMF/EMF OLE previews (MS Equation Editor / MathType)
  - DrawingML images (a:blip)

LaTeX-first strategy (clean KaTeX rendering on the site):
  1. Walk each paragraph; mark OMML + image positions
  2. Convert OMML → LaTeX via pandoc, with pure-Python OMML fallback
  3. Store questions as plain text + $LaTeX$ (not equation PNGs when possible)
  4. Drop image placeholders that sit next to successful OMML→LaTeX
  5. Image → PNG only as last resort when no LaTeX is available
"""
import copy
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import uuid
import zipfile
from concurrent.futures import ThreadPoolExecutor

from django import forms

# ---------------------------------------------------------------------------
# Some older LibreOffice/OpenOffice.org builds write .docx packages using
# pre-standard "purl.oclc.org/ooxml/..." namespace URIs. Rewrite them so
# python-docx can open the file.
# ---------------------------------------------------------------------------
_OOXML_NS_SPECIAL_CASES = [
    (
        "http://purl.oclc.org/ooxml/officeDocument/relationships/extendedProperties",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties",
    ),
    (
        "http://purl.oclc.org/ooxml/officeDocument/extendedProperties",
        "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties",
    ),
]
_OOXML_NS_GENERAL_RE = re.compile(r"http://purl\.oclc\.org/ooxml/([A-Za-z][\w-]*)/")
_OOXML_NS_GENERAL_REPL = r"http://schemas.openxmlformats.org/\1/2006/"


def _rewrite_ooxml_namespace_text(text):
    for old, new in _OOXML_NS_SPECIAL_CASES:
        text = text.replace(old, new)
    return _OOXML_NS_GENERAL_RE.sub(_OOXML_NS_GENERAL_REPL, text)


def normalize_ooxml_namespaces(src_path, dst_path):
    """
    Copies src_path (a .docx package) to dst_path, rewriting old
    'purl.oclc.org/ooxml' namespaces into standard OOXML ones.
    """
    changed = False
    with zipfile.ZipFile(src_path, "r") as zin:
        names = zin.namelist()
        with zipfile.ZipFile(dst_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for name in names:
                data = zin.read(name)
                if name.endswith(".xml") or name.endswith(".rels"):
                    try:
                        text = data.decode("utf-8")
                    except UnicodeDecodeError:
                        zout.writestr(name, data)
                        continue
                    if "purl.oclc.org/ooxml" in text:
                        data = _rewrite_ooxml_namespace_text(text).encode("utf-8")
                        changed = True
                zout.writestr(name, data)
    return changed


try:
    import docx
    from docx.oxml.ns import qn
except ImportError:  # pragma: no cover
    docx = None
    qn = None

from django.core.files.base import ContentFile
from django.core.files.storage import default_storage

# Classic past-paper: (a) (b) (c) (d) … and (1) (2) (3) …
# Supports variable option counts (1 or more), not only A–D.
# Letters a–z and numbers 1–26 are accepted.
_OPTION_TOKEN = r"[a-zA-Z]|[1-9]|1[0-9]|2[0-6]"
OPTION_RE = re.compile(
    rf"(?:^|[\s\u00a0])\(({_OPTION_TOKEN})\)\s*|(?:^|[\s\u00a0])({_OPTION_TOKEN})\)\s+",
)
OPTION_SPLIT_RE = re.compile(
    rf"(?:(?<=\?)|(?<=\s)|(?<=\.))\s*"
    rf"(?:\(({_OPTION_TOKEN})\)|({_OPTION_TOKEN})\))\s*",
    re.IGNORECASE,
)
# Inline / line markers for options
OPTION_MARKER_RE = re.compile(
    rf"(?:\(({_OPTION_TOKEN})\)|({_OPTION_TOKEN})\))\s*",
    re.IGNORECASE,
)
OPTION_LINE_START_RE = re.compile(
    rf"^\s*(?:\(({_OPTION_TOKEN})\)|({_OPTION_TOKEN})\)|([A-Za-z])\.\s*\)|([A-Za-z])[.:])\s*",
    re.IGNORECASE,
)
# GRE / bank style: "A 182" or "B 191" at the start of a line
OPTION_LINE_BARE_RE = re.compile(
    r"^\s*([A-Ia-i])[\t \u00a0]+(\S.*)$",
    re.DOTALL,
)
OPTION_LINE_FULL_RE = re.compile(
    rf"^\s*(?:\(({_OPTION_TOKEN})\)|({_OPTION_TOKEN})\)|([A-Za-z])\.\s*\)|([A-Za-z])[.:])\s*(.*)$",
    re.IGNORECASE | re.DOTALL,
)
# Inline markers including A. / a.) / A: / A 182 (not only (a) / a))
_INLINE_OPT_MARKER_RE = re.compile(
    rf"(?:(?<=^)|(?<=[\s\u00a0])|(?<=[.?!:]))"
    rf"(?:"
    rf"\(({_OPTION_TOKEN})\)|"
    rf"({_OPTION_TOKEN})\)|"
    rf"([A-Ia-i])\.\s*\)|"
    rf"([A-Ia-i])\.|"
    rf"([A-Ia-i]):|"
    rf"([A-Ia-i])[\t \u00a0]+"
    rf")"
    rf"(?=\S)",
    re.IGNORECASE,
)
# "1. …", "1) …", "1 Dinari …", "2 (a) …", "Q.1 …", "Q.12\t…"
QUESTION_NUM_RE = re.compile(
    r"^\s*(?:Q(?:ues(?:tion)?)?[\s.\-:]*)?(\d+)(?:[.)]\s+|\t+|\s+(?=[A-Za-z(]))",
    re.IGNORECASE,
)
ANSWER_HEADING_RE = re.compile(r"ans\w{0,4}r\s*(sheet|key)", re.IGNORECASE)
CORRECT_ANSWER_LINE_RE = re.compile(
    r"^\s*(?:correct\s*answers?|answer\s*key)\s*[:\-–—]\s*(.+?)\s*$",
    re.IGNORECASE,
)
CORRECT_ANSWER_TAIL_RE = re.compile(
    r"(?:^|[\n\r]|[\s\u00a0])(?:correct\s*answers?|answer\s*key)\s*[:\-–—]\s*(.+?)\s*$",
    re.IGNORECASE,
)
MULTI_SELECT_HINT_RE = re.compile(
    r"(?:"
    r"indicate\s+all|select\s+all|choose\s+all|all\s+that\s+apply|"
    r"all\s+such|all\s+correct|"
    r"more\s+than\s+one(?:\s+correct)?(?:\s+(?:option|answer|choice))s?|"
    r"which\s+of\s+the\s+following\s+(?:could|must|are|is/?are)|"
    r"which\s+of\s+the\s+following\s+options\s+are|"
    r"multiple\s+(?:correct|select|answers?)"
    r")",
    re.IGNORECASE,
)
SECTION_HEADING_RE = re.compile(
    r"^\s*section\s*[-–—:]?\s*([IVXLC]+|\d+)\b",
    re.IGNORECASE,
)
PASSAGE_HEAD_RE = re.compile(
    r"^\s*passage\s*\(?\s*Q\.?\s*(\d+)\s*[-–—to]+\s*Q\.?\s*(\d+)",
    re.IGNORECASE,
)
_NUMERIC_ANSWER_RE = re.compile(
    r"^[+\-]?\d+(?:\.\d+)?(?:\s*(?:to|-|–|—)\s*[+\-]?\d+(?:\.\d+)?)?$",
    re.IGNORECASE,
)
_MATCHING_PAIR_RE = re.compile(r"([A-Da-d])\s*[-–—=]\s*([A-Za-z])")

# Map numeric choice labels → a,b,c… (1→a … 26→z)
_OPT_NUM_TO_LETTER = {str(i): chr(ord("a") + i - 1) for i in range(1, 27)}
_LETTER_ORDER = "abcdefghijklmnopqrstuvwxyz"


def _option_letter(token: str) -> str:
    """Normalize '(a)' / '(1)' / 'A' / '12' → lowercase a–z."""
    t = (token or "").strip().lower()
    if t in _OPT_NUM_TO_LETTER:
        return _OPT_NUM_TO_LETTER[t]
    if len(t) == 1 and t in _LETTER_ORDER:
        return t
    return ""


def _options_dict_to_ordered_list(options: dict) -> list[dict]:
    """
    Convert {letter: text} to a stable ordered list of {letter, text}.
    Unknown keys keep insertion order after known a–z keys.
    """
    if not options:
        return []
    ordered = []
    seen = set()
    for ch in _LETTER_ORDER:
        if ch not in options:
            continue
        text = (options.get(ch) or "").strip()
        ordered.append({"letter": ch.upper(), "text": text})
        seen.add(ch)
    for k, text in options.items():
        kk = (k or "").strip().lower()
        if kk in seen:
            continue
        t = (text or "").strip()
        if not t:
            continue
        letter = _option_letter(kk) or kk
        ordered.append({"letter": (letter or kk).upper()[:2], "text": t})
    return ordered


def _marker_letter(match) -> str:
    """Letter from any option-marker regex (first non-empty capturing group)."""
    if not match:
        return ""
    for g in match.groups():
        if g is None:
            continue
        letter = _option_letter(g)
        if letter:
            return letter
    return ""


def _letters_are_sequential(letters: list[str]) -> bool:
    """True when letters are consecutive (A,B or C,D — later rows may start mid-alphabet)."""
    if len(letters) < 2:
        return False
    idxs = []
    for L in letters:
        i = _LETTER_ORDER.find((L or "").lower())
        if i < 0:
            return False
        idxs.append(i)
    return all(b == a + 1 for a, b in zip(idxs, idxs[1:]))


def _normalize_stem_ws(text: str) -> str:
    text = (text or "").replace("\u00a0", " ")
    text = re.sub(r"[\r\n]+", " ", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def _strip_correct_answer(text: str) -> tuple[str, str]:
    """
    Pull a trailing 'Correct Answer : A, D, F' off the block.
    Returns (remaining_text, normalized_answer_letters).
    """
    if not text:
        return "", ""
    m = CORRECT_ANSWER_TAIL_RE.search(text)
    if not m:
        # Whole paragraph is just the answer line
        line = CORRECT_ANSWER_LINE_RE.match(text.replace("\u00a0", " ").strip())
        if line:
            return "", _parse_correct_answer_value(line.group(1))
        return text, ""
    ans = _parse_correct_answer_value(m.group(1))
    remaining = text[: m.start()].rstrip()
    return remaining, ans


def _parse_option_line(line: str) -> tuple[str, str] | None:
    """
    If `line` is a single option, return (letter, text), else None.

    Accepts: (A) 5:2 / A) 5:2 / A. 4 / a.) x / A: 14 / A 182
    Rejects long sentences that happen to start with 'A ' (e.g. 'A zoo has…').
    """
    raw = (line or "").replace("\u00a0", " ").strip()
    if not raw:
        return None
    if CORRECT_ANSWER_LINE_RE.match(raw):
        return None
    m = OPTION_LINE_FULL_RE.match(raw)
    if m:
        # groups: 1=(token) 2=token) 3=letter.) 4=letter.: 5=rest
        letter = _option_letter(
            m.group(1) or m.group(2) or m.group(3) or m.group(4) or ""
        )
        rest = (m.group(5) or "").strip()
        if letter:
            return letter, rest
    m = OPTION_LINE_BARE_RE.match(raw)
    if m:
        letter = _option_letter(m.group(1) or "")
        rest = (m.group(2) or "").strip()
        if not letter or not rest:
            return None
        # Bare "A zoo has twice as many…" is a stem, not an option
        if len(rest) > 90:
            return None
        if re.match(r"^(zoo|the|an?|in|on|of|to|for|if|when|which|what|how)\b", rest, re.I):
            return None
        return letter, rest
    return None


def _options_from_lines(text: str) -> tuple[str, dict]:
    """Split a multi-line question block into stem + {letter: text}."""
    if not text or "\n" not in text:
        return text, {}
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    parsed = []
    for i, line in enumerate(lines):
        parsed.append((i, line, _parse_option_line(line)))

    # Last run of 2+ sequential option lines (blank lines allowed between them)
    best = None
    current = []
    for i, line, opt in parsed:
        if opt is None:
            if (line or "").strip() == "":
                continue
            if current:
                if len(current) >= 2:
                    best = current
                current = []
            continue
        current.append((i, opt[0], opt[1]))
    if len(current) >= 2:
        best = current
    if not best:
        return text, {}

    letters = [item[1] for item in best]
    if not _letters_are_sequential(letters) and len(set(letters)) < 2:
        return text, {}

    first_idx = best[0][0]
    stem_lines = lines[:first_idx]
    options = {}
    for _i, letter, opt_text in best:
        opt_text = re.sub(r"[\s,;]+$", "", (opt_text or "").strip())
        if letter in options and options[letter]:
            options[letter] = f"{options[letter]} {opt_text}".strip()
        else:
            options[letter] = opt_text
    if len(options) < 2:
        return text, {}
    stem = "\n".join(stem_lines).rstrip()
    return stem, options


def _options_from_inline_markers(text: str, marker_re) -> tuple[str, dict]:
    """Group consecutive option markers and take the last multi-option run."""
    if not text:
        return "", {}
    matches = list(marker_re.finditer(text))
    if len(matches) < 1:
        return text, {}

    groups = []
    current = [matches[0]]
    for m in matches[1:]:
        if m.start() - current[-1].end() < 500:
            current.append(m)
        else:
            groups.append(current)
            current = [m]
    groups.append(current)

    candidates = [g for g in groups if len(g) >= 2]
    best = None
    for g in reversed(candidates or []):
        # Prefer the longest sequential suffix (A,B,C…) so a leading
        # "A zoo has…" is not treated as option A.
        trimmed = None
        letters = [_marker_letter(m) for m in g]
        for start in range(len(g)):
            sl = letters[start:]
            if all(sl) and _letters_are_sequential(sl):
                trimmed = g[start:]
                break
        if trimmed and len(trimmed) >= 2:
            best = trimmed
            break
    if not best:
        # Single trailing (a) / (1) marker only
        tail = groups[-1] if groups else []
        if (
            len(tail) == 1
            and tail[0].start() >= max(0, len(text) - 80)
            and marker_re is OPTION_MARKER_RE
        ):
            best = tail
        else:
            return text, {}

    stem = text[: best[0].start()].rstrip()
    stem = re.sub(r"[\s\u00a0]+$", "", stem)
    options = {}
    for i, m in enumerate(best):
        letter = _marker_letter(m) or (
            _LETTER_ORDER[i] if i < len(_LETTER_ORDER) else f"x{i}"
        )
        start = m.end()
        end = best[i + 1].start() if i + 1 < len(best) else len(text)
        opt = text[start:end].strip()
        opt = re.sub(r"[\s,;]+$", "", opt)
        if letter not in options:
            options[letter] = opt
        elif opt:
            options[letter] = f"{options[letter]} {opt}".strip()
    if len(options) < 1:
        return text, {}
    return stem, options


def _options_look_like_mcq(options: dict, stem: str) -> bool:
    """False for past-paper (a)/(b) sub-parts (long text, [marks], blanks)."""
    if not options or len(options) < 2:
        return False
    blob = f"{stem or ''} " + " ".join(options.values())
    if re.search(r"\[\s*\d+\s*\]", blob):
        return False
    if re.search(r"\.{6,}", blob):
        return False
    avg_len = sum(len(v or "") for v in options.values()) / max(len(options), 1)
    if avg_len > 160:
        return False
    return True


def _join_stem_and_options(stem: str, options: dict) -> str:
    parts = [stem.strip()] if (stem or "").strip() else []
    for item in _options_dict_to_ordered_list(options):
        parts.append(f"({item['letter']}) {item['text']}".strip())
    return " ".join(p for p in parts if p).strip()


def _filename_type_hints(name: str) -> list[str]:
    """Question-type keywords found in the uploaded file name."""
    n = re.sub(r"[\s\-]+", "_", (name or "").lower())
    found = []
    checks = (
        (r"multiple_correct|multi_correct|multi_select|more_than_one", "multiple_choice"),
        (r"single_correct|single_choice", "single_choice"),
        (r"true_false|truefalse", "true_false"),
        (r"fill_blank|fill_in", "fill_blank"),
        (r"comprehension|comprehensive", "comprehension"),
        (r"numerical", "numerical"),
        (r"integer", "integer"),
        (r"structured", "structured"),
        (r"matching", "matching"),
    )
    for pat, typ in checks:
        if re.search(pat, n):
            found.append(typ)
    return found


def _parse_correct_answer_value(raw: str) -> str:
    """
    Keep numerical / matching keys intact. Only compress A,B,C letter keys.
    '23' must stay 23 (not letter W).
    """
    raw = re.sub(r"[\s\u00a0]+", " ", (raw or "")).strip()
    if not raw:
        return ""
    pairs = _MATCHING_PAIR_RE.findall(raw)
    if len(pairs) >= 2:
        return ", ".join(f"{a.upper()}-{b.lower()}" for a, b in pairs)
    compact = raw.replace(" ", "")
    if _NUMERIC_ANSWER_RE.match(raw) or _NUMERIC_ANSWER_RE.match(compact):
        return raw
    if re.match(
        r"^[+\-]?\d+(?:\.\d+)?\s+to\s+[+\-]?\d+(?:\.\d+)?$", raw, re.I
    ):
        return raw
    letters = _normalize_answer_token(raw)
    if letters:
        return letters
    return raw


def _looks_numeric_answer(ans: str) -> bool:
    ans = (ans or "").strip()
    if not ans:
        return False
    return bool(
        _NUMERIC_ANSWER_RE.match(ans)
        or _NUMERIC_ANSWER_RE.match(ans.replace(" ", ""))
        or re.match(r"^[+\-]?\d+(?:\.\d+)?\s+to\s+[+\-]?\d+(?:\.\d+)?$", ans, re.I)
    )


def _looks_matching_answer(ans: str) -> bool:
    return len(_MATCHING_PAIR_RE.findall(ans or "")) >= 2


def _looks_comprehension_stem(stem: str) -> bool:
    blob = stem or ""
    if re.search(r"\[\s*\d+\s*\]", blob):
        return True
    if re.search(r"\.{6,}", blob):
        return True
    if re.search(r"\(\s*(?:i{1,3}|iv|v|vi{0,3}|ix|x)\s*\)", blob, re.I):
        return True
    return False


def _infer_section_context(text: str):
    """(question_type, answer_type) from a SECTION / instruction paragraph."""
    t = re.sub(r"\s+", " ", (text or "")).strip().lower()
    if not t:
        return None
    if re.search(r"match(?:ing)?\s+the\s+column|column\s+i\b", t) and re.search(
        r"column|match", t
    ):
        if "match" in t:
            return ("matching", "multiple")
    if re.search(r"\breal number|\bnumerical\b|\binteger type\b", t):
        return ("numerical", "single")
    if re.search(
        r"one or more (?:answers?|options?) are correct|more than one.{0,40}correct",
        t,
    ):
        return ("multiple_choice", "multiple")
    if re.search(
        r"based upon each paragraph|this section contains paragraph|"
        r"passage\s*\(?\s*q",
        t,
    ):
        return ("comprehension", "single")
    if re.search(r"only one is correct|out of which only one", t):
        return ("single_choice", "single")
    if re.search(r"statement\s*-?\s*1", t) and re.search(r"statement\s*-?\s*2", t):
        return ("single_choice", "single")
    return None


def _is_boilerplate_paragraph(text: str) -> bool:
    t = re.sub(r"\s+", " ", (text or "")).strip()
    if not t:
        return False
    if SECTION_HEADING_RE.match(t):
        return True
    if re.match(r"this section contains\b", t, re.I):
        return True
    if PASSAGE_HEAD_RE.match(t):
        return True
    if re.match(r"column\s+(?:i{1,2}|1|2)\b", t, re.I):
        return True
    if re.match(r"each question has 4 choices", t, re.I):
        return True
    if re.match(r"four statements\s*\(", t, re.I):
        return True
    return False


def _closes_current_question(text: str) -> bool:
    """Section / passage headers end the previous question; column titles do not."""
    t = re.sub(r"\s+", " ", (text or "")).strip()
    if not t:
        return False
    if SECTION_HEADING_RE.match(t):
        return True
    if re.match(r"this section contains\b", t, re.I):
        return True
    if PASSAGE_HEAD_RE.match(t):
        return True
    return False


def _force_option_from_paragraph(text: str):
    """Lone option line, including empty / equation-only '(a)' rows."""
    raw = (text or "").replace("\u00a0", " ").strip()
    if not raw:
        return None
    parsed = _parse_option_line(raw)
    if parsed:
        return parsed
    m = OPTION_LINE_START_RE.match(raw)
    if not m:
        m = OPTION_LINE_FULL_RE.match(raw)
    if m and (m.start() == 0):
        letter = _option_letter(
            next((g for g in m.groups() if g), "")
        ) or _marker_letter(m)
        rest = raw[m.end() :].strip()
        if letter:
            return letter, rest
    return None


def _detect_question_and_answer_type(
    stem: str,
    options: dict,
    answer_raw: str,
    filename: str = "",
    section_hint=None,
) -> tuple[str, str]:
    """
    Return (question_type, answer_type).

    Priority: in-document section instructions → answer shape → stem
    wording → options → a single unambiguous file-name hint.
    Mixed file names (single + multiple + numerical) are ignored.
    """
    opts = options or {}
    n_opts = len(opts)
    ans = (answer_raw or "").strip()
    matching_ans = _looks_matching_answer(ans)
    numeric_ans = _looks_numeric_answer(ans)
    letter_parts = []
    if ans and not matching_ans and not numeric_ans:
        letter_parts = [p for p in ans.split(",") if p.strip()]
        # "a,d,f" style; reject long tokens
        if any(len(p.strip()) > 2 for p in letter_parts):
            letter_parts = [
                p for p in _normalize_answer_token(ans).split(",") if p.strip()
            ]
    multi_ans = len(letter_parts) >= 2
    multi_hint = bool(MULTI_SELECT_HINT_RE.search(stem or ""))
    fn_hints = _filename_type_hints(filename)

    if section_hint:
        qtype, atype = section_hint
        if qtype == "single_choice" and multi_ans:
            return "multiple_choice", "multiple"
        if qtype == "numerical":
            return "numerical", "single"
        if qtype == "integer":
            return "integer", "single"
        if qtype == "matching":
            return "matching", "multiple"
        if qtype == "comprehension":
            return "comprehension", "multiple" if (multi_ans or multi_hint) else "single"
        if qtype == "multiple_choice":
            return "multiple_choice", "multiple"
        if qtype == "single_choice" and n_opts >= 2:
            return "single_choice", "single"
        return qtype, atype

    if matching_ans or re.search(r"match(?:ing)?\s+the\s+column", stem or "", re.I):
        return "matching", "multiple"
    if numeric_ans and n_opts < 2:
        return "numerical", "single"

    if n_opts >= 2:
        texts = {(v or "").strip().lower() for v in opts.values() if (v or "").strip()}
        if len(texts) == 2 and texts <= {"true", "false", "t", "f", "yes", "no"}:
            return "true_false", "single"
        if multi_ans or multi_hint:
            return "multiple_choice", "multiple"
        if fn_hints == ["multiple_choice"]:
            return "multiple_choice", "multiple"
        return "single_choice", "single"

    stem_l = stem or ""
    if re.search(r"\b(?:fill\s+in\s+the\s+blank|fills?\s+in)\b", stem_l, re.I):
        return "fill_blank", "single"
    if re.search(
        r"\b(?:integer\s+type|enter\s+(?:the\s+)?(?:correct\s+)?integer)\b",
        stem_l,
        re.I,
    ):
        return "integer", "single"
    if _looks_comprehension_stem(stem_l):
        return "comprehension", "single"
    if len(fn_hints) == 1:
        qtype = fn_hints[0]
        atype = "multiple" if qtype in ("multiple_choice", "matching") else "single"
        return qtype, atype
    return "structured", "single"


def _detect_question_type(stem: str, options: dict, answer_letters: str, filename: str = "") -> str:
    qtype, _atype = _detect_question_and_answer_type(
        stem, options, answer_letters, filename=filename
    )
    return qtype


def _format_correct_answer(ans: str, question_type: str) -> str:
    ans = (ans or "").strip()
    if not ans:
        return ""
    if question_type in ("numerical", "integer", "matching"):
        return ans
    if _looks_numeric_answer(ans) or _looks_matching_answer(ans):
        return ans
    return ans.upper()


def _split_stem_options_answer(text: str) -> tuple[str, dict, str]:
    """
    Split a question block into (stem, options, answer_letters).

    Handles:
      - (A)/(1)/A)/A./a.)/A 182 options on their own lines
      - the same markers inlined on one line
      - a trailing 'Correct Answer : A, D, F' line
    """
    if not text:
        return "", {}, ""
    remaining, answer = _strip_correct_answer(text)
    remaining = remaining.replace("\u00a0", " ")

    stem, options = _options_from_lines(remaining)
    if not options:
        stem, options = _options_from_inline_markers(remaining, OPTION_MARKER_RE)
    if not options:
        stem, options = _options_from_inline_markers(remaining, _INLINE_OPT_MARKER_RE)

    if options and not _options_look_like_mcq(options, stem) and not answer:
        return _normalize_stem_ws(remaining), {}, answer

    return _normalize_stem_ws(stem if options else remaining), options, answer


def _extract_options_from_text(text: str) -> tuple[str, dict]:
    """
    Split stem + options when choices are inline, e.g.:
      '... on p ? (1) 2p (2) 2^{p^2} (3) p2 (4) pp'
      '... is (a) foo (b) bar (c) baz (d) qux'
      '... only (a) single choice'
      '... Indicate all such numbers.\\nA 182\\nB 191'
    Returns (stem, {a:..., b:..., ...}) with 1 or more options.
    """
    stem, options, _ans = _split_stem_options_answer(text)
    return stem, options

VML_NS = "urn:schemas-microsoft-com:vml"
VML_IMAGEDATA_TAG = f"{{{VML_NS}}}imagedata"

MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
OMATH_TAG = f"{{{MATH_NS}}}oMath"
OMATHPARA_TAG = f"{{{MATH_NS}}}oMathPara"

# Placeholders (control chars — never collide with real document text)
_OMATH_PLACEHOLDER_RE = re.compile("\x02OMATH(omatheq\\d+)\x02")
_IMG_PLACEHOLDER_RE = re.compile("\x02IMGRID(rId[^\\x02]+)\x02")

_BATCH_MARKER_RE = re.compile(
    r"OMATHMARKERSTART(omatheq\d+)OMATHMARKEREND\s*\\[\(\[](.*?)\\[\)\]]", re.DOTALL
)

# Equation images are wrapped in <span class="eq-math"> so display height is
# always 1.15em of surrounding text (see courses.equation_display).
def _eq_img_html(url: str) -> str:
    try:
        from courses.equation_display import equation_img_html

        return equation_img_html(url)
    except Exception:
        return (
            f'<img src="{url}" alt="" class="eq-math-img" '
            f'style="height:auto;width:auto;max-height:3.2em;min-height:1.05em;'
            f'vertical-align:middle;margin:0 .2em;background:transparent;border:0;'
            f'object-fit:contain;display:inline-block" '
            f'decoding="async" loading="lazy" />'
        )


# Normalize by MAIN GLYPH height so simple and multi-line formulas share
# similar character size. Display CSS then scales to ~1.3em of body text
# (see templates/partials/equation_inline_css.html) — keep source PNGs modest
# so they do not look huge before JS runs.
_EQ_TARGET_GLYPH_PX = 22   # main text-line height in the PNG (~body text ×1.4)
_EQ_MIN_GLYPH_PX = 14
_EQ_MAX_TOTAL_HEIGHT_PX = 56  # multi-line formulas may be taller overall
_EQ_MAX_WIDTH_PX = 720
_EQ_MIN_HEIGHT_PX = 16  # absolute floor after scale
_EQ_MIN_WIDTH_PX = 16
_CONVERT_BATCH_SIZE = 150
_ENHANCE_WORKERS = min(8, (os.cpu_count() or 4))


# ---------------------------------------------------------------------------
# Paragraph walking — preserve equation order inside the sentence
# ---------------------------------------------------------------------------

def _find_omath_nodes(paragraph_elem):
    return [c for c in paragraph_elem if c.tag in (OMATHPARA_TAG, OMATH_TAG)]


def _rids_from_element(elem):
    """
    Collect image relationship ids from DrawingML blips and VML imagedata.
    Prefer DrawingML (a:blip) when both exist on the same node — they are
    almost always the same equation graphic stored twice.
    """
    if elem is None:
        return []
    blip_rids = []
    for blip in elem.findall(".//" + qn("a:blip")):
        rid = blip.get(qn("r:embed")) or blip.get(
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
        )
        if rid and rid not in blip_rids:
            blip_rids.append(rid)
    if blip_rids:
        return blip_rids
    vml_rids = []
    for imgdata in elem.findall(".//" + VML_IMAGEDATA_TAG):
        rid = imgdata.get(qn("r:id")) or imgdata.get(
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
        )
        if rid and rid not in vml_rids:
            vml_rids.append(rid)
    return vml_rids


def _run_text_with_image_placeholders(run_elem):
    """
    Walk a <w:r> in order so text and embedded equation images keep their
    relative positions (e.g. 'The last digit in [EQ] is').
    """
    parts = []
    for child in run_elem:
        tag = child.tag
        if tag == qn("w:t"):
            parts.append(child.text or "")
        elif tag == qn("w:tab"):
            parts.append("\t")
        elif tag == qn("w:br"):
            # Soft line breaks keep option lines ("A 182" / "A. 4") separate.
            # Page/column breaks are not option separators.
            br_type = child.get(qn("w:type"))
            parts.append(" " if br_type in ("page", "column") else "\n")
        elif tag in (qn("w:drawing"), qn("w:pict"), qn("w:object")):
            for rid in _rids_from_element(child):
                parts.append(f"\x02IMGRID{rid}\x02")
        else:
            # Nested content (e.g. smart tags) — still look for images/text
            for t in child.findall(".//" + qn("w:t")):
                parts.append(t.text or "")
            for rid in _rids_from_element(child):
                parts.append(f"\x02IMGRID{rid}\x02")
    return "".join(parts)


def _paragraph_text_with_placeholders(paragraph, omath_registry):
    """
    Like paragraph.text, but:
      - OMML equations → \x02OMATH{key}\x02
      - DrawingML / VML / OLE equation images → \x02IMGRID{rId}\x02
    so later passes can turn them into LaTeX or inline <img> tags.
    """
    parts = []
    seen_img = set()
    for child in paragraph._p:
        tag = child.tag
        if tag in (OMATH_TAG, OMATHPARA_TAG):
            key = f"omatheq{len(omath_registry)}"
            omath_registry.append((key, copy.deepcopy(child)))
            parts.append(f"\x02OMATH{key}\x02")
        elif tag == qn("w:r"):
            chunk = _run_text_with_image_placeholders(child)
            parts.append(chunk)
            for rid in _IMG_PLACEHOLDER_RE.findall(chunk):
                seen_img.add(rid)
        elif tag == qn("w:hyperlink"):
            for r in child.findall(qn("w:r")):
                parts.append(_run_text_with_image_placeholders(r))
        else:
            # Fallback: any images sitting outside runs
            for rid in _rids_from_element(child):
                if rid not in seen_img:
                    parts.append(f"\x02IMGRID{rid}\x02")
                    seen_img.add(rid)
    text = "".join(parts)
    # Identical back-to-back markers (same rId twice) → keep one
    text = _dedupe_adjacent_img_placeholders(text)
    return text


def _dedupe_adjacent_img_placeholders(text):
    """Collapse consecutive identical IMGRID markers (VML + Drawing pair)."""
    while True:
        new = re.sub(r"(\x02IMGRID(rId[^\x02]+)\x02)\1+", r"\1", text)
        if new == text:
            break
        text = new
    return text


# ---------------------------------------------------------------------------
# OMML → LaTeX (pandoc + pure-Python fallback) — LaTeX-first storage
# ---------------------------------------------------------------------------

def _omml_python_latex(node) -> str:
    try:
        from admin_panel.omml_to_latex import clean_latex, omml_node_to_latex

        return clean_latex(omml_node_to_latex(node))
    except Exception:
        return ""


def _omml_plain_text_fallback(node) -> str:
    texts = []
    for t in node.findall(f".//{{{MATH_NS}}}t"):
        if t.text:
            texts.append(t.text)
    return "".join(texts).strip()


def _render_omath_batch_to_latex(omath_registry, work_dir):
    """Map omatheq keys → LaTeX strings (no $ delimiters)."""
    latex_by_key = {}
    if not omath_registry:
        return latex_by_key

    # 1) Pure-Python OMML conversion for every node (always available)
    for key, node in omath_registry:
        latex = _omml_python_latex(node)
        if latex:
            latex_by_key[key] = latex

    # 2) Prefer pandoc when available (often better for complex OMML)
    if shutil.which("pandoc"):
        try:
            doc = docx.Document()
            for p in list(doc.paragraphs):
                p._p.getparent().remove(p._p)
            for key, node in omath_registry:
                doc.add_paragraph(f"OMATHMARKERSTART{key}OMATHMARKEREND")
                eq_p = doc.add_paragraph()
                eq_p._p.append(copy.deepcopy(node))

            batch_path = os.path.join(work_dir, "omath_batch.docx")
            doc.save(batch_path)
            result = subprocess.run(
                ["pandoc", batch_path, "-t", "latex", "--wrap=none"],
                capture_output=True,
                text=True,
                timeout=90,
                check=True,
            )
            from admin_panel.omml_to_latex import clean_latex

            for key, latex in _BATCH_MARKER_RE.findall(result.stdout or ""):
                latex = clean_latex(latex)
                if latex:
                    # Prefer pandoc when it produced real TeX (has commands)
                    prev = latex_by_key.get(key, "")
                    if "\\" in latex or not prev or len(latex) >= len(prev):
                        latex_by_key[key] = latex
        except (
            subprocess.CalledProcessError,
            subprocess.TimeoutExpired,
            FileNotFoundError,
            OSError,
            Exception,
        ):
            pass

    # 3) Last resort: plain m:t text (still better than a blank)
    for key, node in omath_registry:
        if key not in latex_by_key or not latex_by_key[key]:
            plain = _omml_plain_text_fallback(node)
            if plain:
                latex_by_key[key] = plain

    return latex_by_key


def _drop_images_next_to_omath(text: str, latex_by_key: dict) -> str:
    """When OMML converted to LaTeX, drop adjacent equation image placeholders.

    Word often stores the same formula as OMML + MathType preview image.
    """
    if not text:
        return text

    def repl(m):
        key = m.group(1)
        if latex_by_key.get(key):
            return f"\x02OMATH{key}\x02"
        return m.group(0)

    # OMATH followed by one or more IMGRID
    text = re.sub(
        r"\x02OMATH(omatheq\d+)\x02(?:\s*\x02IMGRIDrId[^\x02]+\x02)+",
        repl,
        text,
    )
    # IMGRID followed by OMATH (rarer order)
    text = re.sub(
        r"(?:\x02IMGRIDrId[^\x02]+\x02\s*)+\x02OMATH(omatheq\d+)\x02",
        repl,
        text,
    )
    return text


def _substitute_placeholders(text, latex_by_key, rid_to_url, *, prefer_latex=True):
    """Replace markers with $LaTeX$ (preferred) or PNG only as last resort."""
    if not text:
        return text

    if prefer_latex:
        text = _drop_images_next_to_omath(text, latex_by_key)

    def omath_repl(m):
        latex = latex_by_key.get(m.group(1))
        if not latex:
            return ""
        from admin_panel.omml_to_latex import wrap_inline_math

        return wrap_inline_math(latex)

    def img_repl(m):
        # LaTeX-first: images only when no alternative
        if not prefer_latex:
            return ""
        rid = m.group(1)
        url = rid_to_url.get(rid)
        if url:
            return _eq_img_html(url)
        return ""

    text = _OMATH_PLACEHOLDER_RE.sub(omath_repl, text)
    text = _IMG_PLACEHOLDER_RE.sub(img_repl, text)
    # Clean whitespace around math
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\s+\$", " $", text)
    text = re.sub(r"\$\s+", "$ ", text)
    text = re.sub(r"\s+(?=<img)", " ", text)
    text = re.sub(r"(?<=>)\s+", " ", text)
    # Drop empty leftover control chars
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
    return text.strip()


# ---------------------------------------------------------------------------
# Document / image conversion helpers
# ---------------------------------------------------------------------------

def _has_numbering(paragraph):
    pPr = paragraph._p.pPr
    return pPr is not None and pPr.numPr is not None


def _soffice_available():
    return _soffice_bin() is not None


def _soffice_bin():
    """
    Locate LibreOffice `soffice` on PATH or common install folders.

    Windows installs rarely put soffice on PATH, so we also probe
    Program Files. Returns absolute path or None.
    """
    for name in ("soffice", "soffice.exe", "libreoffice", "libreoffice.exe"):
        found = shutil.which(name)
        if found and os.path.isfile(found):
            return found

    candidates = []
    # Explicit common paths (Windows + Linux containers)
    candidates.extend(
        [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            "/usr/bin/soffice",
            "/usr/bin/libreoffice",
            "/snap/bin/libreoffice",
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        ]
    )
    # Glob Program Files for versioned installs
    for base in (
        r"C:\Program Files",
        r"C:\Program Files (x86)",
    ):
        try:
            if not os.path.isdir(base):
                continue
            for entry in os.listdir(base):
                if entry.lower().startswith("libreoffice"):
                    candidates.append(
                        os.path.join(base, entry, "program", "soffice.exe")
                    )
        except OSError:
            pass

    for path in candidates:
        if path and os.path.isfile(path):
            return path
    return None


def _soffice_convert_document(src_path, target_ext, outdir):
    """Convert a whole document (e.g. legacy .doc -> .docx) via headless LibreOffice."""
    soffice = _soffice_bin()
    if not soffice:
        return None
    try:
        # User profile dir avoids lock conflicts when multiple conversions run
        profile = os.path.join(outdir, "lo_profile")
        os.makedirs(profile, exist_ok=True)
        # file:/// URL for the profile (LibreOffice expects URI on Windows)
        profile_uri = "file:///" + profile.replace("\\", "/").lstrip("/")
        if not profile_uri.startswith("file:///"):
            profile_uri = "file:///" + profile.replace("\\", "/")
        subprocess.run(
            [
                soffice,
                "--headless",
                "--norestore",
                "--nologo",
                "--nodefault",
                f"-env:UserInstallation={profile_uri}",
                "--convert-to",
                target_ext,
                "--outdir",
                outdir,
                src_path,
            ],
            check=True,
            timeout=180,
            capture_output=True,
        )
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired, FileNotFoundError, OSError):
        return None
    base = os.path.splitext(os.path.basename(src_path))[0]
    out_path = os.path.join(outdir, f"{base}.{target_ext}")
    return out_path if os.path.exists(out_path) and os.path.getsize(out_path) > 0 else None


def _word_com_convert_document(src_path, target_ext, outdir):
    """
    Convert .doc → .docx using installed Microsoft Word (Windows COM).

    Used when LibreOffice is not installed. Preserves equations well enough
    for the OMML → LaTeX pipeline that follows.
    """
    if sys.platform != "win32":
        return None
    if (target_ext or "").lower().lstrip(".") != "docx":
        return None
    if not os.path.isfile(src_path):
        return None

    os.makedirs(outdir, exist_ok=True)
    base = os.path.splitext(os.path.basename(src_path))[0]
    # Avoid overwriting the source when already in outdir
    out_path = os.path.join(outdir, f"{base}_converted.docx")
    if os.path.exists(out_path):
        try:
            os.remove(out_path)
        except OSError:
            pass

    # wdFormatXMLDocument = 16 (.docx)
    # Use absolute paths; Word COM is picky about relative paths.
    src_abs = os.path.abspath(src_path)
    out_abs = os.path.abspath(out_path)
    script_path = os.path.join(outdir, "_word_convert.ps1")
    # Escape single quotes for PowerShell single-quoted strings
    src_ps = src_abs.replace("'", "''")
    out_ps = out_abs.replace("'", "''")
    script = f"""
$ErrorActionPreference = 'Stop'
$src = '{src_ps}'
$dst = '{out_ps}'
$word = $null
$doc = $null
try {{
  $word = New-Object -ComObject Word.Application
  $word.Visible = $false
  $word.DisplayAlerts = 0
  $doc = $word.Documents.Open($src, $false, $true)
  # 16 = wdFormatXMLDocument (.docx)
  $null = $doc.SaveAs([ref]$dst, [ref]16)
  $doc.Close($false) | Out-Null
  $doc = $null
  $word.Quit() | Out-Null
  $word = $null
  if (Test-Path -LiteralPath $dst) {{ Write-Output 'OK' }} else {{ Write-Output 'MISSING' }}
}} catch {{
  Write-Output ('ERR:' + $_.Exception.Message)
  exit 1
}} finally {{
  if ($doc -ne $null) {{ try {{ $doc.Close($false) | Out-Null }} catch {{}} }}
  if ($word -ne $null) {{ try {{ $word.Quit() | Out-Null }} catch {{}} }}
  [System.GC]::Collect()
  [System.GC]::WaitForPendingFinalizers()
}}
"""
    try:
        with open(script_path, "w", encoding="utf-8") as fh:
            fh.write(script)
        result = subprocess.run(
            [
                "powershell",
                "-NoProfile",
                "-ExecutionPolicy",
                "Bypass",
                "-File",
                script_path,
            ],
            check=False,
            timeout=180,
            capture_output=True,
            text=True,
        )
        if result.returncode != 0:
            return None
        if os.path.exists(out_abs) and os.path.getsize(out_abs) > 0:
            return out_abs
    except (subprocess.TimeoutExpired, FileNotFoundError, OSError):
        return None
    return None


def _convert_document_to_docx(src_path, outdir):
    """
    Convert legacy .doc (or broken packages) → .docx.

    Order:
      1. LibreOffice headless (best on Linux / Render)
      2. Microsoft Word COM (Windows desktops without LibreOffice)
    Returns path to .docx or None.
    """
    os.makedirs(outdir, exist_ok=True)
    # Prefer LibreOffice when present
    converted = _soffice_convert_document(src_path, "docx", outdir)
    if converted:
        return converted
    # Fallback: Word on Windows
    converted = _word_com_convert_document(src_path, "docx", outdir)
    if converted:
        return converted
    return None


def _batch_convert_images_soffice(paths, target_ext, outdir):
    mapping = {}
    soffice = _soffice_bin()
    if not paths or not soffice:
        return mapping
    for i in range(0, len(paths), _CONVERT_BATCH_SIZE):
        chunk = paths[i:i + _CONVERT_BATCH_SIZE]
        try:
            subprocess.run(
                [soffice, "--headless", "--norestore", "--convert-to", target_ext,
                 "--outdir", outdir] + chunk,
                check=True, timeout=300, capture_output=True,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired, FileNotFoundError):
            continue
        for p in chunk:
            base = os.path.splitext(os.path.basename(p))[0]
            out_path = os.path.join(outdir, f"{base}.{target_ext}")
            if os.path.exists(out_path):
                mapping[p] = out_path
    return mapping


def _batch_convert_images_windows_gdi(paths, outdir):
    """
    Fast batch WMF/EMF → PNG via a **single** PowerShell process.

    Previous approach launched PowerShell dozens of times and rasterized
    page-sized bitmaps at 3–4× — that made Word uploads hang for minutes.
    """
    mapping = {}
    if not paths or sys.platform != "win32":
        return mapping

    pairs = []
    for src in paths:
        base = os.path.splitext(os.path.basename(src))[0]
        dst = os.path.join(outdir, f"{base}.png")
        if os.path.exists(dst) and os.path.getsize(dst) > 0:
            mapping[src] = dst
            continue
        pairs.append((src, dst))
    if not pairs:
        return mapping

    list_path = os.path.join(outdir, "_wmf_jobs.tsv")
    script_path = os.path.join(outdir, "_wmf_convert.ps1")
    with open(list_path, "w", encoding="utf-8") as fh:
        for src, dst in pairs:
            fh.write(f"{src}\t{dst}\n")

    script = r'''
Add-Type -AssemblyName System.Drawing
$ErrorActionPreference = 'Continue'
$listPath = $args[0]
$ok = 0
Get-Content -LiteralPath $listPath -Encoding UTF8 | ForEach-Object {
  if (-not $_) { return }
  $parts = $_.Split("`t")
  if ($parts.Count -lt 2) { return }
  $src = $parts[0]
  $dst = $parts[1]
  try {
    $srcImg = [System.Drawing.Image]::FromFile($src)
    $nw = [Math]::Max(1, $srcImg.Width)
    $nh = [Math]::Max(1, $srcImg.Height)
    $scale = 2
    if ($nw -le 80 -and $nh -le 80) { $scale = 4 }
    elseif ($nw -gt 350 -or $nh -gt 350) { $scale = 1 }
    $w = [Math]::Max(1, [int]($nw * $scale))
    $h = [Math]::Max(1, [int]($nh * $scale))
    $maxSide = 900
    if ($w -gt $maxSide -or $h -gt $maxSide) {
      $r = $maxSide / [Math]::Max($w, $h)
      $w = [Math]::Max(1, [int]($w * $r))
      $h = [Math]::Max(1, [int]($h * $r))
    }
    $bmp = New-Object System.Drawing.Bitmap $w, $h
    $bmp.SetResolution(144, 144)
    $g = [System.Drawing.Graphics]::FromImage($bmp)
    $g.Clear([System.Drawing.Color]::White)
    $g.InterpolationMode = [System.Drawing.Drawing2D.InterpolationMode]::HighQualityBicubic
    $g.SmoothingMode = [System.Drawing.Drawing2D.SmoothingMode]::HighQuality
    $g.PixelOffsetMode = [System.Drawing.Drawing2D.PixelOffsetMode]::HighQuality
    $g.DrawImage($srcImg, 0, 0, $w, $h)
    $bmp.Save($dst, [System.Drawing.Imaging.ImageFormat]::Png)
    $g.Dispose(); $bmp.Dispose(); $srcImg.Dispose()
    $ok++
  } catch { }
}
Write-Output $ok
'''
    with open(script_path, "w", encoding="utf-8") as fh:
        fh.write(script)

    try:
        timeout = min(480, max(60, 30 + int(len(pairs) * 0.15)))
        subprocess.run(
            [
                "powershell",
                "-NoProfile",
                "-ExecutionPolicy",
                "Bypass",
                "-File",
                script_path,
                list_path,
            ],
            check=False,
            timeout=timeout,
            capture_output=True,
            text=True,
        )
    except (subprocess.TimeoutExpired, FileNotFoundError, OSError):
        pass

    for src, dst in pairs:
        if os.path.exists(dst) and os.path.getsize(dst) > 0:
            mapping[src] = dst
    return mapping


def _batch_convert_images(paths, target_ext, outdir):
    """Convert many WMF/EMF images to PNG (deduped paths, single PS batch)."""
    if not paths:
        return {}
    unique = list(dict.fromkeys(paths))
    mapping = _batch_convert_images_soffice(unique, target_ext, outdir)
    remaining = [p for p in unique if p not in mapping]
    if remaining and target_ext.lower() == "png":
        mapping.update(_batch_convert_images_windows_gdi(remaining, outdir))
    return mapping


def _enhance_equation_pngs_parallel(paths):
    """Crop/resize many PNGs in parallel."""
    paths = [p for p in dict.fromkeys(paths) if p and os.path.exists(p)]
    if not paths:
        return
    if len(paths) == 1:
        _enhance_equation_png(paths[0])
        return
    workers = min(_ENHANCE_WORKERS, len(paths))
    with ThreadPoolExecutor(max_workers=workers) as pool:
        list(pool.map(_enhance_equation_png, paths))


def _estimate_main_glyph_height(gray_img):
    """
    Estimate the height of the primary text line from row ink density.

    Simple single-line formulas → almost full content height.
    Formulas with fractions/exponents → longest dense band (main baseline).
    Scaling so this height is constant makes 'log y = cx' and dense options
    share the same optical character size.
    """
    w, h = gray_img.size
    if w < 2 or h < 2:
        return max(h, 1)

    # Row ink counts (pixels darker than near-white)
    rows = []
    pix = gray_img.load()
    for y in range(h):
        ink = 0
        for x in range(w):
            if pix[x, y] < 230:
                ink += 1
        rows.append(ink)

    max_ink = max(rows) if rows else 0
    if max_ink < 1:
        return max(h, 1)

    thr = max(2, int(max_ink * 0.12))
    active = [i for i, v in enumerate(rows) if v >= thr]
    if not active:
        return max(h, 1)

    # Contiguous bands of active rows
    bands = []
    start = active[0]
    prev = active[0]
    for i in active[1:]:
        if i == prev + 1:
            prev = i
            continue
        bands.append((start, prev))
        start = prev = i
    bands.append((start, prev))

    band_heights = [b[1] - b[0] + 1 for b in bands]
    longest = max(band_heights)
    total_span = active[-1] - active[0] + 1

    # Multi-line / fraction: use longest band as "main glyph" height
    if total_span > longest * 1.55 and len(bands) >= 2:
        return max(longest, 6)
    # Single visual line (maybe with exponents still one band)
    return max(total_span, 6)


def _enhance_equation_png(path, padding=2):
    """
    Crop + scale so the MAIN GLYPH line is always ~28px tall.

    Then display with height:auto — simple and complex formulas share the
    same character size (no more huge log(...) vs tiny fractions).
    """
    try:
        from PIL import Image, ImageEnhance, ImageOps, ImageFilter
    except ImportError:
        return
    try:
        im = Image.open(path)
        if im.mode in ("RGBA", "LA") or (im.mode == "P" and "transparency" in im.info):
            bg = Image.new("RGBA", im.size, (255, 255, 255, 255))
            im = Image.alpha_composite(bg, im.convert("RGBA")).convert("RGB")
        else:
            im = im.convert("RGB")

        gray = im.convert("L")
        mask = gray.point(lambda p: 255 if p < 248 else 0)
        mask = mask.filter(ImageFilter.MaxFilter(3))
        bbox = mask.getbbox()
        if bbox:
            left, top, right, bottom = bbox
            left = max(0, left - padding)
            top = max(0, top - padding)
            right = min(im.width, right + padding)
            bottom = min(im.height, bottom + padding)
            im = im.crop((left, top, right, bottom))
            gray = im.convert("L")

        w, h = im.size
        if w < 2 or h < 2:
            return

        glyph_h = _estimate_main_glyph_height(gray)
        glyph_h = max(glyph_h, 4)

        scale = _EQ_TARGET_GLYPH_PX / float(glyph_h)
        # Clamp insane scales
        scale = min(max(scale, 0.25), 10.0)

        new_w = max(1, int(round(w * scale)))
        new_h = max(1, int(round(h * scale)))

        # Cap overall dimensions
        if new_h > _EQ_MAX_TOTAL_HEIGHT_PX:
            s2 = _EQ_MAX_TOTAL_HEIGHT_PX / float(new_h)
            new_w = max(1, int(round(new_w * s2)))
            new_h = _EQ_MAX_TOTAL_HEIGHT_PX
        if new_w > _EQ_MAX_WIDTH_PX:
            s2 = _EQ_MAX_WIDTH_PX / float(new_w)
            new_h = max(1, int(round(new_h * s2)))
            new_w = _EQ_MAX_WIDTH_PX
        if new_h < _EQ_MIN_HEIGHT_PX:
            s2 = _EQ_MIN_HEIGHT_PX / float(max(new_h, 1))
            new_w = max(1, int(round(new_w * s2)))
            new_h = _EQ_MIN_HEIGHT_PX

        if (new_w, new_h) != (w, h):
            im = im.resize((new_w, new_h), Image.Resampling.LANCZOS)

        im = ImageEnhance.Contrast(im).enhance(1.12)
        im = ImageEnhance.Sharpness(im).enhance(1.2)
        # Minimal white pad only — large borders made equations look like boxes
        im = ImageOps.expand(im, border=1, fill=(255, 255, 255))
        im.save(path, format="PNG", optimize=True)
    except Exception:
        pass


def _autocrop_png(path, padding=2):
    """Back-compat name used by the import pipeline."""
    _enhance_equation_png(path, padding=padding)

def _save_image_blob(blob, ext, media_subdir):
    """Save raw image bytes to MEDIA; return public URL."""
    ext = (ext or "png").lstrip(".").lower()
    if ext in ("wmf", "emf"):
        # Should have been converted already — store as-is only as last resort
        pass
    fname = f"{media_subdir}/{uuid.uuid4().hex}.{ext}"
    saved_name = default_storage.save(fname, ContentFile(blob))
    try:
        return default_storage.url(saved_name)
    except Exception:
        return saved_name


_ANSWER_LETTER_RE = re.compile(r"([a-zA-Z])")


def _normalize_answer_token(ans: str) -> str:
    """Turn 'c', 'C.', '2', ' a,b ' into clean lowercase letter key(s)."""
    if not ans:
        return ""
    ans = ans.strip().lower()
    # Numeric keys from (1)(2)(3)… style papers
    if re.fullmatch(r"[1-9]|1[0-9]|2[0-6]", ans):
        return _OPT_NUM_TO_LETTER.get(ans, "")
    # multi-select e.g. "a,b" or "1,2"
    parts = re.split(r"[\s,;/]+", ans)
    seen = []
    for p in parts:
        p = p.strip().strip("().")
        if not p:
            continue
        if p in _OPT_NUM_TO_LETTER:
            L = _OPT_NUM_TO_LETTER[p]
        elif len(p) == 1 and p in _LETTER_ORDER:
            L = p
        else:
            m = _ANSWER_LETTER_RE.search(p)
            L = m.group(1).lower() if m else ""
            if L in _OPT_NUM_TO_LETTER:
                L = _OPT_NUM_TO_LETTER[L]
        if L and L not in seen:
            seen.append(L)
    return ",".join(seen)


def _fix_answer_number_row(cells):
    """
    Force sequential numbering on answer-sheet number rows.

    Word tables often typo the last cell of a block (e.g. 61–69 then 61
    instead of 70). For a standard 1..10 grid we renumber from the first
    digit so keys line up with the answer row beneath.
    """
    vals = [c.strip() for c in cells]
    digit_idxs = [i for i, v in enumerate(vals) if v.isdigit()]
    if len(digit_idxs) < 2:
        return vals
    start = int(vals[digit_idxs[0]])
    for j, i in enumerate(digit_idxs):
        vals[i] = str(start + j)
    return vals


def _is_two_column_qa_table(rows) -> bool:
    """
    Detect simple answer-key tables produced by convert_mcq_docs.py:

        Question | Answer
        1        | B
        2        | C
        ...
    """
    if not rows or len(rows) < 2:
        return False
    # Inspect up to first 8 data rows for "digit | letter" pairs
    digit_letter = 0
    checked = 0
    for row in rows[:12]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 2:
            continue
        left, right = cells[0], cells[1]
        # Skip header
        if left.lower() in {"question", "q.no", "q", "no", "number", "s.no", "s.no."}:
            continue
        if right.lower() in {"answer", "ans", "key", "correct"}:
            continue
        checked += 1
        left_ok = bool(re.fullmatch(r"\d+", left))
        right_ok = bool(_normalize_answer_token(right))
        if left_ok and right_ok:
            digit_letter += 1
    return checked > 0 and digit_letter >= max(1, checked // 2)


def _parse_two_column_qa_table(rows) -> dict:
    """Parse Question|Answer rows into {question_number: letter}."""
    answers = {}
    for row in rows:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 2:
            continue
        left, right = cells[0], cells[1]
        if not left.isdigit():
            continue
        letter = _normalize_answer_token(right)
        if not letter:
            continue
        # Prefer first mapping if duplicates
        if left not in answers:
            answers[left] = letter
    return answers


def _looks_like_number_header_row(cells) -> bool:
    """True when a majority of non-empty cells are question numbers."""
    vals = [c.strip() for c in cells if c.strip()]
    if len(vals) < 2:
        return False
    digits = sum(1 for v in vals if v.isdigit())
    return digits >= max(2, (len(vals) + 1) // 2)


def _looks_like_answer_letter_row(cells) -> bool:
    """True when cells look like answer keys (a/b/c… or 1/2/3…), not Q numbers."""
    vals = [c.strip() for c in cells if c.strip()]
    if len(vals) < 2:
        return False
    letterish = 0
    for v in vals:
        if _normalize_answer_token(v):
            letterish += 1
        elif v.isdigit() and int(v) <= 26:
            # bare numbers can be (1)(2) style keys — count as letterish
            letterish += 1
    return letterish >= max(2, (len(vals) + 1) // 2)


def _parse_answer_sheet_tables(document):
    """
    Parse Answer Sheet tables in either format:

    1) Two-column key (convert_mcq_docs / latex_ready):
        Question | Answer
        1        | B
        2        | C

    2) Classic past-paper grid:
        1  2  3  … 10
        c  d  c  … b
        11 12 … 20
        d  a  … b
    """
    answers = {}
    for table in document.tables:
        rows = table.rows
        if not rows:
            continue

        # ---- Format 1: two-column Question | Answer ----
        if _is_two_column_qa_table(rows):
            parsed = _parse_two_column_qa_table(rows)
            for num, letter in parsed.items():
                if num not in answers:
                    answers[num] = letter
            continue

        # ---- Format 2: classic number-row / answer-row pairs ----
        i = 0
        while i + 1 < len(rows):
            raw_nums = [c.text.strip() for c in rows[i].cells]
            raw_ans = [c.text.strip() for c in rows[i + 1].cells]

            if not _looks_like_number_header_row(raw_nums):
                i += 1
                continue
            # Next row must look like answers, not another number sequence
            # (avoids pairing 1|B with next row 2|C as "1→2")
            if not _looks_like_answer_letter_row(raw_ans):
                i += 1
                continue

            num_cells = _fix_answer_number_row(raw_nums)
            ans_cells = raw_ans
            for num, ans in zip(num_cells, ans_cells):
                if not num.isdigit():
                    continue
                letter = _normalize_answer_token(ans)
                if not letter:
                    continue
                if num in answers and answers[num] != letter:
                    continue
                answers[num] = letter
            i += 2
    return answers


def _normalize_display_latex(text: str) -> str:
    """
    Convert common LaTeX delimiters from convert_mcq_docs output to $...$
    so MathJax / KaTeX on the site render consistently.
      \\( x \\)  →  $x$
      \\[ x \\]  →  $$x$$
    """
    if not text:
        return text
    # Display math first
    text = re.sub(r"\\\[(.+?)\\\]", r"$$\1$$", text, flags=re.DOTALL)
    # Inline math
    text = re.sub(r"\\\((.+?)\\\)", r"$\1$", text, flags=re.DOTALL)
    return text


# ---------------------------------------------------------------------------
# Main entry
# ---------------------------------------------------------------------------

def parse_docx_questions(uploaded_file, media_subdir="question_equations"):
    """
    Parse an uploaded .doc/.docx and return row-dicts with:
      - question_text / option_a..d as plain text + $LaTeX$ (KaTeX)
      - equation_images: only leftover image URLs when LaTeX was unavailable
    """
    if docx is None:
        raise forms.ValidationError(
            "The 'python-docx' package isn't installed on the server. "
            "Run: pip install python-docx"
        )

    work_dir = tempfile.mkdtemp(prefix="docx_import_")
    try:
        src_name = uploaded_file.name
        src_path = os.path.join(work_dir, os.path.basename(src_name))
        with open(src_path, "wb") as f:
            for chunk in uploaded_file.chunks():
                f.write(chunk)

        docx_path = src_path
        name_lower = (src_name or "").lower()
        # Treat only true legacy binary .doc (not .docx)
        is_legacy_doc = name_lower.endswith(".doc") and not name_lower.endswith(".docx")

        if is_legacy_doc:
            convert_dir = os.path.join(work_dir, "converted")
            os.makedirs(convert_dir, exist_ok=True)
            converted = _convert_document_to_docx(src_path, convert_dir)
            if not converted:
                lo = "found" if _soffice_bin() else "not found"
                on_render = bool(
                    os.environ.get("RENDER")
                    or os.environ.get("RENDER_EXTERNAL_HOSTNAME")
                )
                if on_render and lo == "not found":
                    raise forms.ValidationError(
                        "This server cannot convert legacy .doc files yet "
                        "(LibreOffice is not installed on the live host). "
                        "Quick fix: open the file in Word → Save As → .docx, "
                        "then upload the .docx. "
                        "Permanent fix: redeploy with the Docker image that "
                        "includes LibreOffice (see Dockerfile / render.yaml)."
                    )
                word_hint = (
                    " Microsoft Word is also unavailable for COM conversion."
                    if sys.platform == "win32"
                    else " On Linux/Render, install LibreOffice (soffice) or upload .docx."
                )
                raise forms.ValidationError(
                    "Couldn't convert this legacy .doc file to .docx "
                    f"(LibreOffice {lo}.{word_hint}) "
                    "Install LibreOffice, or open the file in Word → Save As → "
                    ".docx and re-upload. On Windows PCs, Microsoft Word COM is used "
                    "automatically when LibreOffice is missing."
                )
            docx_path = converted

        normalized_path = os.path.join(
            work_dir, "normalized_" + os.path.basename(docx_path)
        )
        try:
            normalize_ooxml_namespaces(docx_path, normalized_path)
            docx_path = normalized_path
        except (zipfile.BadZipFile, OSError):
            pass

        try:
            document = docx.Document(docx_path)
        except Exception as first_error:
            # Re-save through Word/LibreOffice to repair odd packages
            renorm_dir = os.path.join(work_dir, "renormalized")
            os.makedirs(renorm_dir, exist_ok=True)
            renormalized = _convert_document_to_docx(docx_path, renorm_dir)
            if renormalized:
                try:
                    # Normalize namespaces again after conversion
                    fixed = os.path.join(renorm_dir, "normalized_retry.docx")
                    try:
                        normalize_ooxml_namespaces(renormalized, fixed)
                        renormalized = fixed
                    except (zipfile.BadZipFile, OSError):
                        pass
                    document = docx.Document(renormalized)
                    docx_path = renormalized
                except Exception:
                    raise forms.ValidationError(
                        f"Couldn't read this Word file after conversion: {first_error}"
                    )
            else:
                raise forms.ValidationError(
                    f"Couldn't read this Word file: {first_error}"
                )

        # ---- 1. Collect every image part (WMF/EMF/PNG/JPG/…) ----
        # De-dupe by content hash so identical equation blobs convert once.
        image_dir = os.path.join(work_dir, "equations")
        os.makedirs(image_dir, exist_ok=True)

        rid_to_src = {}          # rid → local source path (any format)
        rid_to_ext = {}
        need_convert = []       # unique paths that are wmf/emf
        blob_hash_to_path = {}  # md5 → path (dedupe)

        for rel_id, rel in document.part.rels.items():
            if "image" not in rel.reltype:
                continue
            try:
                target = rel.target_part
                ext = (target.partname.ext or "").lower().lstrip(".")
                blob = target.blob
            except Exception:
                continue
            if not blob:
                continue
            digest = hashlib.md5(blob).hexdigest()
            if digest in blob_hash_to_path:
                src_img_path = blob_hash_to_path[digest]
            else:
                src_img_path = os.path.join(
                    image_dir, f"{digest}.{ext or 'bin'}"
                )
                with open(src_img_path, "wb") as f:
                    f.write(blob)
                blob_hash_to_path[digest] = src_img_path
                if ext in ("wmf", "emf"):
                    need_convert.append(src_img_path)
            rid_to_src[rel_id] = src_img_path
            rid_to_ext[rel_id] = ext or "bin"

        converted_map = _batch_convert_images(need_convert, "png", image_dir)
        # Parallel crop/resize (major speedup vs sequential PIL)
        _enhance_equation_pngs_parallel(list(converted_map.values()))

        # rid → final local PNG/image path ready to upload
        rid_to_local = {}
        for rid, src in rid_to_src.items():
            if src in converted_map:
                rid_to_local[rid] = converted_map[src]
            elif rid_to_ext.get(rid) in ("png", "jpg", "jpeg", "gif", "webp", "bmp"):
                rid_to_local[rid] = src
            # else: unconverted wmf left out (will skip that marker)

        # ---- 2. Walk paragraphs → questions (with inline placeholders) ----
        questions = []
        current = None
        omath_registry = []
        section_hint = None  # (question_type, answer_type) from SECTION text
        current_passage = ""
        pending_shared_options = {}
        last_pending_letter = ""

        def start_question(num_hint=None):
            nonlocal current
            passage = ""
            if section_hint and section_hint[0] == "comprehension" and current_passage:
                passage = current_passage
            current = {
                "question_text": "",
                "options": {},
                "image_rids": [],
                "question_number": num_hint or str(len(questions) + 1),
                "inline_answer": "",
                "section_hint": section_hint,
                "passage": passage,
            }
            questions.append(current)

        def _absorb_block(block_text: str):
            """Merge a paragraph (stem / options / Correct Answer) into current."""
            if current is None:
                return
            stem, opts, ans = _split_stem_options_answer(block_text)
            if not opts:
                forced = _force_option_from_paragraph(block_text)
                if forced:
                    opts = {forced[0]: forced[1]}
                    stem = ""
            if ans and not current.get("inline_answer"):
                current["inline_answer"] = ans
            if opts:
                if stem:
                    current["question_text"] = (
                        f"{current['question_text']} {stem}".strip()
                        if current["question_text"]
                        else stem
                    )
                for letter, opt_text in opts.items():
                    if letter in current["options"] and current["options"][letter]:
                        if opt_text:
                            current["options"][letter] = (
                                f"{current['options'][letter]} {opt_text}".strip()
                            )
                    else:
                        current["options"][letter] = opt_text
                return
            if stem:
                current["question_text"] = (
                    f"{current['question_text']} {stem}".strip()
                    if current["question_text"]
                    else stem
                )

        for p in document.paragraphs:
            raw_text = _paragraph_text_with_placeholders(p, omath_registry).strip()

            visible_plain = re.sub(
                r"\x02(?:OMATH|IMGRID)[^\x02]+\x02", "", raw_text
            )
            visible_plain = visible_plain.replace("\u00a0", " ").strip()
            # "Correct Answer : A, D" is the key for the current question,
            # not a new "Answer Sheet" section heading.
            if (
                ANSWER_HEADING_RE.search(visible_plain)
                and not CORRECT_ANSWER_LINE_RE.match(visible_plain)
                and not current
            ):
                break
            if (
                ANSWER_HEADING_RE.search(visible_plain)
                and re.search(r"sheet", visible_plain, re.I)
                and not CORRECT_ANSWER_LINE_RE.match(visible_plain)
            ):
                break

            plain_for_num = _IMG_PLACEHOLDER_RE.sub("", _OMATH_PLACEHOLDER_RE.sub("", raw_text))
            explicit_num = QUESTION_NUM_RE.match(plain_for_num)
            is_new_q = _has_numbering(p) or bool(explicit_num)

            rids_here = _IMG_PLACEHOLDER_RE.findall(raw_text)

            if not raw_text and not rids_here:
                continue

            if _is_boilerplate_paragraph(visible_plain):
                # Section / passage headers end the previous question.
                # Column titles are skipped but must not drop the current Q
                # (its Correct Answer line still follows).
                if _closes_current_question(visible_plain):
                    current = None
                    inferred = _infer_section_context(visible_plain)
                    if inferred:
                        section_hint = inferred
                        if inferred[0] != "comprehension":
                            current_passage = ""
                        pending_shared_options = {}
                        last_pending_letter = ""
                    p_head = PASSAGE_HEAD_RE.match(visible_plain)
                    if p_head:
                        current_passage = ""
                        section_hint = section_hint or ("comprehension", "single")
                continue

            text = raw_text
            if explicit_num:
                # Strip leading "1." / "Q.12" from the plain prefix only
                text = QUESTION_NUM_RE.sub("", text, count=1)

            if is_new_q:
                num_hint = explicit_num.group(1) if explicit_num else None
                start_question(num_hint)
                _absorb_block(text)
                numerical_section = bool(
                    section_hint and section_hint[0] in ("numerical", "integer")
                )
                if (
                    not current["options"]
                    and pending_shared_options
                    and not numerical_section
                ):
                    current["options"] = dict(pending_shared_options)
                if numerical_section:
                    # Intervals like (a, b) are not MCQ options
                    current["options"] = {}
                current["image_rids"].extend(rids_here)
                continue

            if current is None:
                forced = _force_option_from_paragraph(text)
                if forced:
                    pending_shared_options[forced[0]] = forced[1]
                    last_pending_letter = forced[0]
                    continue
                if last_pending_letter:
                    extra = _normalize_stem_ws(text)
                    if extra:
                        prev = pending_shared_options.get(last_pending_letter) or ""
                        pending_shared_options[last_pending_letter] = (
                            f"{prev} {extra}".strip()
                        )
                    continue
                if section_hint and section_hint[0] == "comprehension":
                    extra = _normalize_stem_ws(text)
                    if extra:
                        current_passage = f"{current_passage} {extra}".strip()
                continue

            _absorb_block(text)
            current["image_rids"].extend(rids_here)

        # ---- 2b. Pull inline options / Correct Answer out of the stem ----
        for q in questions:
            leftover = q.get("question_text") or ""
            stem, opts, ans = _split_stem_options_answer(leftover)
            if ans and not q.get("inline_answer"):
                q["inline_answer"] = ans
            if opts and not q.get("options"):
                q["question_text"] = stem
                q["options"] = opts
            elif ans:
                q["question_text"] = stem
            if q.get("options") and not _options_look_like_mcq(
                q["options"], q.get("question_text") or ""
            ) and not q.get("inline_answer"):
                q["question_text"] = _join_stem_and_options(
                    q.get("question_text") or "", q["options"]
                )
                q["options"] = {}

        # ---- 3. Answer Sheet table ----
        answers = _parse_answer_sheet_tables(document)

        # ---- 3b. OMML → LaTeX ----
        latex_by_key = _render_omath_batch_to_latex(omath_registry, work_dir)

        # ---- 4. Upload images & build rid → URL map (used by all questions) ----
        # Same local file (content-hash dedupe) → one MEDIA write, shared URL.
        rid_to_url = {}
        path_to_url = {}
        for rid, local_path in rid_to_local.items():
            if local_path in path_to_url:
                rid_to_url[rid] = path_to_url[local_path]
                continue
            ext = os.path.splitext(local_path)[1].lstrip(".") or "png"
            try:
                with open(local_path, "rb") as fh:
                    blob = fh.read()
                if not blob:
                    continue
                url = _save_image_blob(blob, ext, media_subdir)
                path_to_url[local_path] = url
                rid_to_url[rid] = url
            except OSError:
                continue

        # ---- 5. Assemble output rows with equations inlined ----
        rows_out = []
        for position, q in enumerate(questions, start=1):
            qnum = q["question_number"] or str(position)

            # Unique equation URLs for this question (review UI + equation_images field)
            saved_urls = []
            seen = set()
            for rid in q["image_rids"]:
                url = rid_to_url.get(rid)
                if url and url not in seen:
                    seen.add(url)
                    saved_urls.append(url)
            # Also pick up any rids only present as placeholders in text
            for rid in _IMG_PLACEHOLDER_RE.findall(
                q["question_text"] + " " + " ".join(q["options"].values())
            ):
                url = rid_to_url.get(rid)
                if url and url not in seen:
                    seen.add(url)
                    saved_urls.append(url)

            options = q["options"] or {}
            # Match answer key by question_number; fall back to 1-based position
            # or an in-document "Correct Answer : A, D" / "45.92" line.
            table_ans = answers.get(str(qnum), "") or answers.get(str(position), "")
            if table_ans:
                answer_raw = _normalize_answer_token(table_ans)
            else:
                answer_raw = q.get("inline_answer") or ""
            question_type, answer_type = _detect_question_and_answer_type(
                q.get("question_text") or "",
                options,
                answer_raw,
                filename=src_name,
                section_hint=q.get("section_hint"),
            )
            if question_type in ("numerical", "integer"):
                options = {}

            q_text = _normalize_display_latex(
                _substitute_placeholders(
                    q["question_text"].strip(),
                    latex_by_key,
                    rid_to_url,
                    prefer_latex=True,
                )
            )

            # Build full options list (any count ≥ 1) with LaTeX/images inlined
            options_list = []
            for item in _options_dict_to_ordered_list(options):
                rendered = _normalize_display_latex(
                    _substitute_placeholders(
                        item.get("text") or "",
                        latex_by_key,
                        rid_to_url,
                        prefer_latex=True,
                    )
                )
                options_list.append(
                    {
                        "letter": item["letter"],
                        "text": rendered,
                    }
                )

            # Legacy A–D columns (first four options) for older UI/export paths
            by_letter = {o["letter"].lower(): o["text"] for o in options_list}
            opt_a = by_letter.get("a", "")
            opt_b = by_letter.get("b", "")
            opt_c = by_letter.get("c", "")
            opt_d = by_letter.get("d", "")
            # If options used non a–d keys only, still fill A–D by order
            if options_list and not any([opt_a, opt_b, opt_c, opt_d]):
                for i, o in enumerate(options_list[:4]):
                    key = _LETTER_ORDER[i]
                    if key == "a":
                        opt_a = o["text"]
                    elif key == "b":
                        opt_b = o["text"]
                    elif key == "c":
                        opt_c = o["text"]
                    elif key == "d":
                        opt_d = o["text"]

            expl = ""

            # Only keep equation image URLs that are still referenced as <img>
            combined = " ".join(
                [q_text, opt_a, opt_b, opt_c, opt_d]
                + [o["text"] for o in options_list]
            )
            still_used = [
                u for u in saved_urls if u and u in combined
            ]
            # Last resort: if we have images but zero math made it into text
            if still_used == [] and saved_urls and "$" not in combined and "<img" not in combined:
                # Prefer not to dump raw images into stem when LaTeX path failed
                # for everything — keep list for admin review only
                pass

            passage_text = _normalize_display_latex(
                _substitute_placeholders(
                    (q.get("passage") or "").strip(),
                    latex_by_key,
                    rid_to_url,
                    prefer_latex=True,
                )
            ) if q.get("passage") else ""

            row = {
                "question_text": q_text,
                "question_type": question_type,
                "answer_type": answer_type,
                "passage": passage_text,
                "option_a": opt_a,
                "option_b": opt_b,
                "option_c": opt_c,
                "option_d": opt_d,
                # Full list of options (1..N) for import → QuestionOption rows
                "options_list": options_list,
                "correct_answer": _format_correct_answer(answer_raw, question_type),
                "marks": 1,
                "explanation": expl,
                "topic": "",
                "paper_code": "",
                "year": None,
                "season": "",
                "zone": "",
                "question_number": qnum,
                # Review only: images not embedded when LaTeX succeeded
                "equation_images": still_used,
            }
            rows_out.append(row)

        if not rows_out:
            raise forms.ValidationError(
                "No questions were found in this Word file. Expected numbered "
                "questions (1. …) with options like (a) (b) (c) (d) or (1) (2)…"
            )

        return rows_out
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)
